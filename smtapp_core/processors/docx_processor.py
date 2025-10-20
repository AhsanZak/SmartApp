"""
DOCX file processor
"""
from typing import Dict, Any
from docx import Document
import os

from processors.base_processor import BaseProcessor


class DOCXProcessor(BaseProcessor):
    """Processor for DOCX files"""
    
    def supports(self, file_extension: str) -> bool:
        """Check if DOCX is supported"""
        return file_extension.lower() in ["docx", "doc"]
    
    def process(self, file_path: str) -> Dict[str, Any]:
        """
        Process DOCX file and extract text
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            # Get metadata
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "file_size": os.path.getsize(file_path)
            }
            
            # Add core properties if available
            if doc.core_properties:
                metadata.update({
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                    "subject": doc.core_properties.subject or "",
                    "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                })
            
            return {
                "text": text.strip(),
                "metadata": metadata
            }
            
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")

